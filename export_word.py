from docx import Document
from docx.shared import Inches
import pandas as pd

def export_to_word(filename, title, raw_data, results, figures):
    doc = Document()
    doc.add_heading(title, level=1)

    # Сирі дані
    doc.add_heading("Сирі дані", level=2)
    table = doc.add_table(rows=1, cols=len(raw_data.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(raw_data.columns):
        hdr_cells[i].text = col
    for _, row in raw_data.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    # Результати аналізу
    doc.add_heading("Результати аналізу", level=2)
    doc.add_paragraph(str(results))

    # Додавання графіків
    doc.add_heading("Графіки", level=2)
    for fig_path in figures:
        doc.add_picture(fig_path, width=Inches(5))

    doc.save(filename)
