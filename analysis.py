# analysis.py
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
import statsmodels.api as sm
from statsmodels.formula.api import ols
import docx

def run_analysis(df):
    # Перевірка нормальності (Shapiro–Wilk)
    shapiro_results = {col: stats.shapiro(df[col].dropna()) for col in df.select_dtypes(include='number').columns}

    # Побудова графіків
    plt.figure(figsize=(8, 6))
    sns.boxplot(data=df.select_dtypes(include='number'))
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.savefig("boxplot.png")
    plt.close()

    # Однофакторний приклад ANOVA
    if df.shape[1] >= 2:
        formula = f'{df.columns[0]} ~ C({df.columns[1]})'
        model = ols(formula, data=df).fit()
        anova_table = sm.stats.anova_lm(model, typ=2)
    else:
        anova_table = None

    # Збереження у Word
    doc = docx.Document()
    doc.add_heading("Результати статистичного аналізу", 0)
    doc.add_paragraph("Перевірка нормальності (Shapiro–Wilk):")
    for col, res in shapiro_results.items():
        doc.add_paragraph(f"{col}: W={res[0]:.4f}, p={res[1]:.4f}")
    if anova_table is not None:
        doc.add_paragraph("ANOVA:")
        doc.add_paragraph(str(anova_table))
    doc.add_picture("boxplot.png", width=docx.shared.Inches(5))
    output_path = "result.docx"
    doc.save(output_path)

    return output_path
