import sys
import pandas as pd
import numpy as np
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTableWidget, QTableWidgetItem, QFileDialog, QLabel,
    QMessageBox, QComboBox, QHeaderView
)
from PyQt5.QtGui import QIcon
from scipy import stats
import matplotlib.pyplot as plt
from docx import Document

class StatystykaApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Статистика — Чаплоуцький А.М., УНУ")
        self.setGeometry(200, 200, 800, 600)
        self.setWindowIcon(QIcon("icon.ico"))

        self.data = None

        layout = QVBoxLayout()

        # Вибір типу аналізу
        self.analysis_type = QComboBox()
        self.analysis_type.addItems([
            "Однофакторний",
            "Двофакторний",
            "Трьохфакторний",
            "Повторні вимірювання"
        ])
        layout.addWidget(QLabel("Оберіть тип аналізу:"))
        layout.addWidget(self.analysis_type)

        # Таблиця введення даних
        self.table = QTableWidget(5, 5)
        self.table.setHorizontalHeaderLabels(["Колонка 1", "Колонка 2", "Колонка 3", "Колонка 4", "Колонка 5"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        layout.addWidget(self.table)

        # Кнопки
        btn_layout = QHBoxLayout()
        load_btn = QPushButton("Завантажити з Excel")
        load_btn.clicked.connect(self.load_from_excel)
        btn_layout.addWidget(load_btn)

        save_btn = QPushButton("Зберегти у Word")
        save_btn.clicked.connect(self.save_to_word)
        btn_layout.addWidget(save_btn)

        analyze_btn = QPushButton("Виконати аналіз")
        analyze_btn.clicked.connect(self.run_analysis)
        btn_layout.addWidget(analyze_btn)

        plot_btn = QPushButton("Побудувати графік")
        plot_btn.clicked.connect(self.plot_data)
        btn_layout.addWidget(plot_btn)

        layout.addLayout(btn_layout)

        container = QWidget()
        container.setLayout(layout)
        self.setCentralWidget(container)

    def load_from_excel(self):
        path, _ = QFileDialog.getOpenFileName(self, "Відкрити Excel", "", "Excel Files (*.xlsx *.xls)")
        if path:
            df = pd.read_excel(path)
            self.table.setRowCount(len(df))
            self.table.setColumnCount(len(df.columns))
            self.table.setHorizontalHeaderLabels(df.columns)
            for i in range(len(df)):
                for j in range(len(df.columns)):
                    self.table.setItem(i, j, QTableWidgetItem(str(df.iat[i, j])))

    def save_to_word(self):
        path, _ = QFileDialog.getSaveFileName(self, "Зберегти у Word", "", "Word Files (*.docx)")
        if path:
            doc = Document()
            doc.add_heading("Результати аналізу", level=1)
            if self.data is not None:
                for i, row in self.data.iterrows():
                    doc.add_paragraph(str(list(row)))
            doc.save(path)
            QMessageBox.information(self, "Готово", "Файл збережено успішно!")

    def run_analysis(self):
        # Зчитуємо дані з таблиці
        rows = self.table.rowCount()
        cols = self.table.columnCount()
        data_list = []
        for i in range(rows):
            row_data = []
            for j in range(cols):
                item = self.table.item(i, j)
                if item and item.text():
                    try:
                        row_data.append(float(item.text()))
                    except ValueError:
                        row_data.append(np.nan)
                else:
                    row_data.append(np.nan)
            data_list.append(row_data)

        df = pd.DataFrame(data_list).dropna(how='all')
        self.data = df

        # Перевірка нормальності
        shapiro_pvalues = []
        for col in df.columns:
            col_data = df[col].dropna()
            if len(col_data) > 3:
                _, p_value = stats.shapiro(col_data)
                shapiro_pvalues.append(p_value)

        if any(p < 0.05 for p in shapiro_pvalues):
            QMessageBox.warning(self, "Увага", "Дані не відповідають нормальному розподілу (p < 0.05)!")
        else:
            QMessageBox.information(self, "OK", "Дані відповідають нормальному розподілу.")

    def plot_data(self):
        if self.data is not None:
            self.data.plot(kind="bar")
            plt.show()
        else:
            QMessageBox.warning(self, "Помилка", "Немає даних для побудови графіка.")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = StatystykaApp()
    window.show()
    sys.exit(app.exec_())
